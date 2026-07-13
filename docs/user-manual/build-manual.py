#!/usr/bin/env python3
"""Build Activity Lead operator manual as DOCX and PDF from markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "activity-lead-operator-manual.md"
DOCX_OUT = ROOT / "Activity-Lead-Operator-Manual.docx"
PDF_OUT = ROOT / "Activity-Lead-Operator-Manual.pdf"
IMAGE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
MISSING_IMAGES: list[str] = []


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Calibri"
        heading.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text.strip())


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text.strip())


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    img_path = (ROOT / rel_path).resolve()
    if not img_path.is_file():
        MISSING_IMAGES.append(rel_path)
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot pending: {alt or rel_path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        return

    doc.add_picture(str(img_path), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alt.strip():
        caption = doc.add_paragraph(alt.strip())
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption.runs:
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(9)
            caption.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def markdown_to_docx(md_text: str) -> Document:
    doc = Document()
    set_document_defaults(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                table.rows[0].cells[col].text = header
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    if c_idx < len(table.rows[r_idx + 1].cells):
                        table.rows[r_idx + 1].cells[c_idx].text = cell
            doc.add_paragraph()
            continue

        if re.match(r"^[-*]\s+", stripped):
            add_bullet(doc, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith("---"):
            i += 1
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        # Bold inline **text** simplified
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1

    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return
    except Exception as exc:  # noqa: BLE001
        print(f"docx2pdf failed: {exc}", file=sys.stderr)

    try:
        import comtypes.client  # type: ignore

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        return
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "PDF export requires Microsoft Word. Open the DOCX file and use File → Save as PDF."
        ) from exc


def main() -> int:
    global MISSING_IMAGES
    MISSING_IMAGES = []

    if not SOURCE.exists():
        print(f"Missing source: {SOURCE}", file=sys.stderr)
        return 1

    md_text = SOURCE.read_text(encoding="utf-8")
    doc = markdown_to_docx(md_text)
    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")

    if MISSING_IMAGES:
        print(
            f"Note: {len(MISSING_IMAGES)} screenshot(s) not found yet — "
            "see docs/user-manual/screenshots/README.md",
            file=sys.stderr,
        )

    try:
        export_pdf(DOCX_OUT, PDF_OUT)
        print(f"Wrote {PDF_OUT}")
    except RuntimeError as err:
        print(str(err), file=sys.stderr)
        return 2

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
